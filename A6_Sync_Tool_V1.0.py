"""
UNFCCC Article 6.2 Sync Tool
============================

Reads Technical Expert Review Reports (TERRs) in .docx format, extracts
findings / CB needs / metadata, writes them into the Excel database
(tbl_TER_Status, tbl_Metadata_Findings, tbl_CB_Needs), and rebuilds the
Dashboard sheet.

Supports both Format A (ID-first tables) and Format B (requirement-first
tables, including the B2 variant with Significant/Persistent columns).

FDR parsing (Finding / During / Recommendation) uses a regex classifier
for simple cases and optionally calls the Anthropic API for complex ones
(numbered blocks, multiple During sections, etc.). Falls back to regex
when ANTHROPIC_API_KEY is not set or the call fails.

Entry points:
    sync_all(docx_paths, xlsx_path, log)      -- full sync + dashboard
    refresh_dashboard_only(xlsx_path, log)    -- rebuild dashboard only

Requires: python-docx, openpyxl
"""

from __future__ import annotations

import json
import os
import re
import sys
import threading
import tkinter as tk
import urllib.request
from datetime import date
from tkinter import filedialog, messagebox
from typing import Any

# ---------------------------------------------------------------------------
# Dependency check (runs before heavier imports so users get a clear prompt)
# ---------------------------------------------------------------------------

_MISSING: list[str] = []
try:
    from docx import Document
except ImportError:
    _MISSING.append("python-docx")

try:
    from openpyxl import load_workbook
    from openpyxl.chart import BarChart, PieChart, Reference
    from openpyxl.chart.series import DataPoint, SeriesLabel
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    _MISSING.append("openpyxl")

if _MISSING:
    import subprocess

    _root = tk.Tk()
    _root.withdraw()
    if messagebox.askyesno(
        "Setup required",
        f"The following packages need to be installed: {', '.join(_MISSING)}\n\nProceed?",
    ):
        subprocess.check_call([sys.executable, "-m", "pip", "install", *_MISSING])
        messagebox.showinfo("Setup complete", "Please restart the application.")
    sys.exit()


# ---------------------------------------------------------------------------
# UNFCCC / UN brand palette
# ---------------------------------------------------------------------------

PALETTE = {
    "navy":        "00205B",
    "blue":        "009EDB",
    "blue_light":  "D6EEF8",
    "gold":        "C9A84C",
    "gold_light":  "FDF3DC",
    "green":       "3A7D44",
    "green_light": "D9EFD7",
    "red":         "C0392B",
    "red_light":   "FAD7D3",
    "grey":        "F2F4F6",
    "white":       "FFFFFF",
    "text_dark":   "1A1A2E",
    "text_mid":    "4A5568",
}

# Short alias used throughout the dashboard builder
C = PALETTE


# ---------------------------------------------------------------------------
# Low-level styling helpers for openpyxl
# ---------------------------------------------------------------------------

def _fill(hex_color: str) -> PatternFill:
    return PatternFill("solid", fgColor=hex_color)


def _font(bold=False, color="1A1A2E", size=11, italic=False, name="Calibri") -> Font:
    return Font(bold=bold, color=color, size=size, italic=italic, name=name)


def _align(h="left", v="center", wrap=False) -> Alignment:
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)


def _border_thin(sides: str = "bottom") -> Border:
    thin = Side(style="thin", color="BDBDBD")
    return Border(**{side: thin for side in sides.split(",")})


def _border_full() -> Border:
    thin = Side(style="thin", color="BDBDBD")
    return Border(left=thin, right=thin, top=thin, bottom=thin)


def _set(
    cell, value="", bold=False, fg=None, font_color="1A1A2E",
    h_align="left", v_align="center", size=11, wrap=False,
    border=None, italic=False,
):
    """Apply value + styling to a single cell in one call."""
    cell.value = value
    cell.font = _font(bold=bold, color=font_color, size=size, italic=italic)
    if fg:
        cell.fill = _fill(fg)
    cell.alignment = _align(h_align, v_align, wrap)
    if border:
        cell.border = border


def _merge_set(ws, r1, c1, r2, c2, **kwargs):
    """Merge a rectangular range and style the anchor cell."""
    ws.merge_cells(start_row=r1, start_column=c1, end_row=r2, end_column=c2)
    _set(ws.cell(r1, c1), **kwargs)


def _row_height(ws, row, height):
    ws.row_dimensions[row].height = height


def _col_width(ws, col, width):
    ws.column_dimensions[get_column_letter(col)].width = width


# ---------------------------------------------------------------------------
# Generic sheet-reading helpers (used by the dashboard aggregator)
# ---------------------------------------------------------------------------

def _headers(ws) -> dict[str, int]:
    """Return {header_name: column_index} for row 1 of the sheet."""
    headers = {}
    for col in range(1, ws.max_column + 1):
        value = ws.cell(1, col).value
        if value and value not in headers:
            headers[str(value).strip()] = col
    return headers


def _read_sheet(ws) -> list[dict[str, Any]]:
    """Read a sheet into a list of {header: value} dictionaries."""
    headers = _headers(ws)
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            rows.append({
                name: (row[idx - 1] if idx <= len(row) else None)
                for name, idx in headers.items()
            })
    return rows


def _get(row: dict, *keys, default: str = "") -> str:
    """Return the first non-empty value among the given keys, as a stripped string."""
    for key in keys:
        value = row.get(key)
        if value is not None:
            return str(value).strip()
    return default


# ---------------------------------------------------------------------------
# Dashboard data aggregation
# ---------------------------------------------------------------------------

def collect_dashboard_data(wb) -> dict[str, Any]:
    """
    Aggregate all metrics needed to build the dashboard sheet.

    Safely handles empty or missing sheets by returning zeros / empty lists.
    """
    ter_rows = _read_sheet(wb["tbl_TER_Status"])        if "tbl_TER_Status"        in wb.sheetnames else []
    fn_rows  = _read_sheet(wb["tbl_Metadata_Findings"]) if "tbl_Metadata_Findings" in wb.sheetnames else []
    cb_rows  = _read_sheet(wb["tbl_CB_Needs"])          if "tbl_CB_Needs"          in wb.sheetnames else []

    # ---- Headline KPIs ----------------------------------------------------
    parties      = sorted({_get(r, "Party_Name", "Party_Code") for r in ter_rows if _get(r, "Party_Name", "Party_Code")})
    review_years = sorted({_get(r, "Review_Year") for r in ter_rows if _get(r, "Review_Year")})
    n_parties    = len(parties)
    n_ters       = len(ter_rows)
    n_findings   = len(fn_rows)
    n_cb         = len(cb_rows)

    # ---- Recommendation compliance: flagged vs. actually filled in -------
    rec_flagged = sum(1 for r in fn_rows if _get(r, "Recommendation_Flag") == "Yes")
    rec_filled  = sum(
        1 for r in fn_rows
        if _get(r, "Recommendation_Flag") == "Yes"
        and len(_get(r, "Recommendation_Text")) > 10
    )
    rec_rate = round(rec_filled / rec_flagged * 100, 1) if rec_flagged else 0

    # ---- FDR completeness: all three parts have real content -------------
    fdr_complete = sum(
        1 for r in fn_rows
        if len(_get(r, "Finding_Text")) > 5
        and len(_get(r, "During_Text")) > 5
        and len(_get(r, "Recommendation_Text")) > 5
    )
    fdr_rate = round(fdr_complete / n_findings * 100, 1) if n_findings else 0

    # ---- Parties grouped by Report Type ----------------------------------
    rtype_parties: dict[str, set] = {}
    for r in ter_rows:
        rt = _get(r, "Report_Type") or "Unknown"
        p  = _get(r, "Party_Name", "Party_Code") or "Unknown"
        rtype_parties.setdefault(rt, set()).add(p)
    rtype_party_counts = {rt: len(ps) for rt, ps in rtype_parties.items()}

    # ---- Submission status by Paragraph Group ----------------------------
    parties_with_ter = {_get(r, "Party_Name", "Party_Code") for r in ter_rows}
    group_parties_with: dict[str, set] = {}
    group_parties_all: dict[str, set]  = {}
    for r in fn_rows:
        pg = _get(r, "Paragraph_Group", "Para_Group") or "Unknown"
        p  = _get(r, "Party_Name", "Party_Code") or "Unknown"
        group_parties_all.setdefault(pg, set()).add(p)
        if p in parties_with_ter:
            group_parties_with.setdefault(pg, set()).add(p)

    group_submission_table = []
    for pg in sorted(group_parties_all):
        total   = len(group_parties_all.get(pg, set()))
        with_s  = len(group_parties_with.get(pg, set()))
        without = total - with_s
        group_submission_table.append((pg, with_s, without, total))

    # ---- Findings by Paragraph Group -------------------------------------
    para_counts: dict[str, int] = {}
    for r in fn_rows:
        pg = _get(r, "Paragraph_Group", "Para_Group") or "Unknown"
        para_counts[pg] = para_counts.get(pg, 0) + 1
    para_table = sorted(para_counts.items())

    # ---- Per-party findings and CB breakdown -----------------------------
    party_fn: dict[str, int] = {}
    party_cb: dict[str, int] = {}
    for r in fn_rows:
        p = _get(r, "Party_Name", "Party_Code") or "Unknown"
        party_fn[p] = party_fn.get(p, 0) + 1
    for r in cb_rows:
        p = _get(r, "Party_Name", "Party_Code") or "Unknown"
        party_cb[p] = party_cb.get(p, 0) + 1
    all_parties = sorted(set(party_fn) | set(party_cb))
    party_table = [(p, party_fn.get(p, 0), party_cb.get(p, 0)) for p in all_parties]

    # ---- Significance split (Significant / Minor / Unclassified) ---------
    sig_counts = {"Significant": 0, "Minor": 0, "Unclassified": 0}
    for r in fn_rows:
        s = _get(r, "Significance") or "Unclassified"
        if s not in sig_counts:
            s = "Unclassified"
        sig_counts[s] += 1

    # ---- Issue-type breakdown --------------------------------------------
    issue_counts: dict[str, int] = {}
    for r in fn_rows:
        it = _get(r, "Issue_Type") or "Not classified"
        issue_counts[it] = issue_counts.get(it, 0) + 1
    issue_table = sorted(issue_counts.items(), key=lambda x: -x[1])

    # ---- Top 10 requirements by finding count ----------------------------
    req_counts: dict[str, int] = {}
    req_labels: dict[str, str] = {}
    for r in fn_rows:
        rid = _get(r, "Requirement_ID") or "Unknown"
        lbl = _get(r, "Paragraph_Title") or rid
        req_counts[rid] = req_counts.get(rid, 0) + 1
        req_labels[rid] = lbl
    top10_req = sorted(req_counts.items(), key=lambda x: -x[1])[:10]
    top10_req = [(rid, req_labels[rid], cnt) for rid, cnt in top10_req]

    # ---- CB needs heat-map (party x paragraph group) ---------------------
    cb_heatmap: dict[str, dict[str, int]] = {}
    para_groups_cb: set[str] = set()
    for r in cb_rows:
        p  = _get(r, "Party_Name", "Party_Code") or "Unknown"
        pg = _get(r, "Paragraph_Group") or "Unknown"
        para_groups_cb.add(pg)
        cb_heatmap.setdefault(p, {})
        cb_heatmap[p][pg] = cb_heatmap[p].get(pg, 0) + 1

    # ---- Report type counts ----------------------------------------------
    rtype_counts: dict[str, int] = {}
    for r in ter_rows:
        rt = _get(r, "Report_Type") or "Unknown"
        rtype_counts[rt] = rtype_counts.get(rt, 0) + 1

    # ---- Findings trend by Review Year -----------------------------------
    year_counts: dict[str, int] = {}
    for r in fn_rows:
        y = _get(r, "Review_Year") or "Unknown"
        year_counts[y] = year_counts.get(y, 0) + 1
    year_trend = sorted(year_counts.items())

    # ---- Requirement coverage matrix (party x top requirements) ----------
    top_req_ids = [rid for rid, _, _ in top10_req]
    party_req_matrix: dict[str, dict[str, int]] = {}
    for r in fn_rows:
        p   = _get(r, "Party_Name", "Party_Code") or "Unknown"
        rid = _get(r, "Requirement_ID") or "Unknown"
        if rid in top_req_ids:
            party_req_matrix.setdefault(p, {})
            party_req_matrix[p][rid] = party_req_matrix[p].get(rid, 0) + 1

    narrative = _build_narrative(
        n_parties, n_ters, n_findings, n_cb,
        sig_counts, rtype_party_counts, review_years,
    )

    return {
        "parties": parties,
        "review_years": review_years,
        "n_parties": n_parties,
        "n_ters": n_ters,
        "n_findings": n_findings,
        "n_cb": n_cb,
        "rec_flagged": rec_flagged,
        "rec_filled": rec_filled,
        "rec_rate": rec_rate,
        "fdr_rate": fdr_rate,
        "para_table": para_table,
        "party_table": party_table,
        "sig_counts": sig_counts,
        "issue_table": issue_table,
        "top10_req": top10_req,
        "cb_heatmap": cb_heatmap,
        "para_groups_cb": sorted(para_groups_cb),
        "rtype_counts": rtype_counts,
        "rtype_party_counts": rtype_party_counts,
        "group_submission_table": group_submission_table,
        "year_trend": year_trend,
        "top_req_ids": top_req_ids,
        "party_req_matrix": party_req_matrix,
        "matrix_parties": sorted(party_req_matrix.keys()),
        "narrative": narrative,
    }


def _build_narrative(
    n_parties, n_ters, n_findings, n_cb,
    sig_counts, rtype_party_counts, review_years,
) -> str:
    """
    Produce a short narrative summary for the dashboard insight card.

    Calls the Anthropic API if ANTHROPIC_API_KEY is set; otherwise returns a
    deterministic template-based summary so the dashboard always renders.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")

    fallback = (
        f"As of this refresh, {n_parties} Parties have been reviewed across "
        f"{n_ters} TER report(s). A total of {n_findings} inconsistencies were "
        f"identified, of which {sig_counts['Significant']} are significant and "
        f"{sig_counts['Minor']} minor. {n_cb} capacity-building needs were recorded."
    )

    if not api_key or n_parties == 0:
        return fallback

    summary = (
        f"Parties reviewed: {n_parties}. "
        f"TER reports: {n_ters}. "
        f"Total inconsistencies: {n_findings} "
        f"(Significant: {sig_counts['Significant']}, Minor: {sig_counts['Minor']}). "
        f"CB needs: {n_cb}. "
        f"Report types: {', '.join(f'{k}={v}' for k, v in rtype_party_counts.items())}. "
        f"Review years: {', '.join(review_years) if review_years else 'N/A'}."
    )

    try:
        payload = json.dumps({
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 200,
            "system": (
                "You write concise 2-sentence dashboard narrative insights for "
                "UNFCCC Article 6.2 Technical Expert Review dashboards. "
                "Be factual, professional, and highlight the most notable numbers. "
                "No bullet points. No markdown. Plain text only."
            ),
            "messages": [{
                "role": "user",
                "content": f"Write a 2-sentence dashboard narrative insight from this data:\n{summary}",
            }],
        }).encode()

        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = json.loads(resp.read())
        return data["content"][0]["text"].strip()
    except Exception:
        return fallback


# ---------------------------------------------------------------------------
# Dashboard builder
# ---------------------------------------------------------------------------

def build_dashboard(wb, log=None):
    """
    Rebuild the 'Dashboard' sheet from scratch using the current workbook data.

    Layout:
        1.  Header banner + last-refreshed date
        2.  KPI row (Parties, TERs, Findings, CB Needs, Significant, Rec Rate, FDR Rate)
        3.  Calculated insight card (narrative)
        4.  Parties by record type + significance split
        5.  Findings trend by review year
        6.  Parties with / without findings by paragraph group (stacked)
        7.  Inconsistencies by paragraph group
        8.  Severity-scored party review outcomes
        9.  Issue type breakdown
        10. Top-10 requirements leaderboard
        11. Requirement coverage matrix (party x top requirements)
        12. CB needs heat-map
        13. Report type distribution
        14. Footer
    """
    def _log(msg):
        if log:
            log(msg)

    _log("\nBuilding Dashboard sheet...")

    if "Dashboard" in wb.sheetnames:
        del wb["Dashboard"]

    ws = wb.create_sheet("Dashboard", 0)
    ws.sheet_view.showGridLines = False
    ws.sheet_view.zoomScale = 85

    data = collect_dashboard_data(wb)

    # Column widths
    for col, width in {1: 2, 2: 24, 3: 15, 4: 15, 5: 15, 6: 15, 7: 2,
                       8: 24, 9: 13, 10: 13, 11: 13, 12: 2}.items():
        _col_width(ws, col, width)

    row = 1

    # --- 1. Header banner --------------------------------------------------
    for r in range(1, 6):
        for c in range(1, 13):
            ws.cell(r, c).fill = _fill(C["navy"])
    _row_height(ws, 1, 6)
    _row_height(ws, 2, 38)
    _row_height(ws, 3, 22)
    _row_height(ws, 4, 16)
    _row_height(ws, 5, 6)

    _merge_set(ws, 2, 2, 2, 11,
               value="ARTICLE 6.2  \u2014  TECHNICAL EXPERT REVIEW DASHBOARD",
               bold=True, fg=C["navy"], font_color=C["gold"], h_align="left", size=20)

    year_suffix = "  |  Years: " + ", ".join(data["review_years"]) if data["review_years"] else ""
    _merge_set(ws, 3, 2, 3, 11,
               value=f"  UNFCCC Transparency Framework  \u2014  Cooperative Approaches{year_suffix}",
               fg=C["navy"], font_color="A0B4CC", h_align="left", size=10)
    _merge_set(ws, 4, 2, 4, 11,
               value=f"  Last refreshed: {date.today().strftime('%d %B %Y')}",
               fg=C["navy"], font_color="506070", h_align="left", size=9, italic=True)

    row = 6

    # --- 2. KPI cards ------------------------------------------------------
    _row_height(ws, row, 8)
    row += 1

    avg_fn = round(data["n_findings"] / data["n_parties"], 1) if data["n_parties"] else 0
    kpis = [
        ("Parties",      data["n_parties"],                  "Under review",           C["blue_light"],  C["blue"]),
        ("TER Reports",  data["n_ters"],                     "Total records",          C["gold_light"],  C["gold"]),
        ("Findings",     data["n_findings"],                 f"Avg {avg_fn}/party",    C["red_light"],   C["red"]),
        ("CB Needs",     data["n_cb"],                       "Identified",             C["green_light"], C["green"]),
        ("Significant",  data["sig_counts"]["Significant"],  "Flagged significant",    "FAD7D3",         C["red"]),
        ("Rec Rate",     f"{data['rec_rate']}%",             "Recommendations filled", "FDF3DC",         C["gold"]),
        ("FDR Complete", f"{data['fdr_rate']}%",             "Full F+D+R rows",        "D6EEF8",         C["blue"]),
    ]

    top_cards    = kpis[:4]
    bottom_cards = kpis[4:]
    top_cols     = [2, 4, 6, 9]
    bottom_cols  = [2, 5, 9]

    def _kpi_card(r_start, c_start, title, value, sub, bg, value_color, span=2):
        for rr in range(r_start, r_start + 5):
            for cc in range(c_start, c_start + span):
                ws.cell(rr, cc).fill = _fill(bg)
        _row_height(ws, r_start,     4)
        _row_height(ws, r_start + 1, 36)
        _row_height(ws, r_start + 2, 18)
        _row_height(ws, r_start + 3, 14)
        _row_height(ws, r_start + 4, 6)
        _merge_set(ws, r_start + 1, c_start, r_start + 1, c_start + span - 1,
                   value=str(value), bold=True, fg=bg, font_color=value_color,
                   h_align="center", size=26)
        _merge_set(ws, r_start + 2, c_start, r_start + 2, c_start + span - 1,
                   value=title, bold=True, fg=bg, font_color="2D3748",
                   h_align="center", size=10)
        _merge_set(ws, r_start + 3, c_start, r_start + 3, c_start + span - 1,
                   value=sub, fg=bg, font_color="718096",
                   h_align="center", size=8, italic=True)

    for (title, value, sub, bg, vc), col in zip(top_cards, top_cols):
        _kpi_card(row, col, title, value, sub, bg, vc, span=2)
    for (title, value, sub, bg, vc), col in zip(bottom_cards, bottom_cols):
        _kpi_card(row + 5, col, title, value, sub, bg, vc, span=2)

    row += 11
    _row_height(ws, row, 10)
    row += 1

    # --- Section header helpers ------------------------------------------
    def _section_header(r, c1, c2, title, color=C["navy"]):
        _merge_set(ws, r, c1, r, c2, value=f"  {title}",
                   bold=True, fg=color, font_color=C["white"],
                   h_align="left", size=11)
        _row_height(ws, r, 22)

    # --- 3. Calculated insight card ---------------------------------------
    _section_header(row, 2, 11, "\u25fc  CALCULATED INSIGHT", color=C["gold"])
    row += 1
    _row_height(ws, row, 5)
    row += 1
    _merge_set(ws, row, 2, row + 3, 11,
               value=data.get("narrative", ""),
               fg=C["gold_light"], font_color=C["text_dark"],
               h_align="left", v_align="top", wrap=True, size=10, italic=True)
    for r in range(row, row + 4):
        _row_height(ws, r, 18)
    row += 4
    _row_height(ws, row, 10)
    row += 1

    # --- 4. Parties by record type + significance split -------------------
    _section_header(row, 2, 6, "\u25fc  PARTIES BY RECORD TYPE", color=C["blue"])
    _section_header(row, 8, 11, "\u25fc  SIGNIFICANCE SPLIT", color=C["red"])
    row += 1
    _row_height(ws, row, 5)
    row += 1

    rtype_hdr = row
    for col, (txt, align) in zip([2, 3, 4], [("Record Type", "left"), ("Parties", "center"), ("TERs", "center")]):
        _set(ws.cell(row, col), value=txt, bold=True, fg=C["blue"],
             font_color=C["white"], h_align=align, size=10)
    for col, (txt, align) in zip([8, 9, 10], [("Significance", "left"), ("Count", "center"), ("Share", "center")]):
        _set(ws.cell(row, col), value=txt, bold=True, fg=C["red"],
             font_color=C["white"], h_align=align, size=10)
    _row_height(ws, row, 18)
    row += 1

    rtype_start = row
    sorted_rt = sorted(data["rtype_party_counts"].items(), key=lambda x: -x[1])
    for i, (rtype, pc) in enumerate(sorted_rt):
        bg = C["grey"] if i % 2 == 0 else C["white"]
        tc = data["rtype_counts"].get(rtype, 0)
        _set(ws.cell(row, 2), value=rtype, fg=bg, h_align="left", size=10)
        _set(ws.cell(row, 3), value=pc, fg=bg, h_align="center", size=10, bold=True, font_color=C["blue"])
        _set(ws.cell(row, 4), value=tc, fg=bg, h_align="center", size=10)
        _row_height(ws, row, 17)
        row += 1
    rtype_end = row - 1

    # Significance panel (same rows on the right)
    sig_start   = rtype_start
    sig_labels  = ["Significant", "Minor", "Unclassified"]
    sig_colors  = [C["red"], C["gold"], C["text_mid"]]
    total_fn    = data["n_findings"] or 1
    for i, (lbl, clr) in enumerate(zip(sig_labels, sig_colors)):
        cnt = data["sig_counts"].get(lbl, 0)
        pct = f"{cnt / total_fn * 100:.1f}%"
        bg  = C["grey"] if i % 2 == 0 else C["white"]
        r2  = sig_start + i
        _set(ws.cell(r2, 8),  value=lbl, fg=bg, size=10)
        _set(ws.cell(r2, 9),  value=cnt, fg=bg, h_align="center", size=10, bold=True, font_color=clr)
        _set(ws.cell(r2, 10), value=pct, fg=bg, h_align="center", size=10)
        _row_height(ws, r2, 17)
    sig_end = sig_start + len(sig_labels) - 1

    if sorted_rt:
        ch = BarChart()
        ch.type = "bar"
        ch.style = 10
        ch.title = "Parties by Record Type"
        ch.y_axis.title = "Record Type"
        ch.x_axis.title = "Parties"
        ch.width = 13
        ch.height = max(7, len(sorted_rt) * 1.4)
        ch.add_data(Reference(ws, min_col=3, min_row=rtype_start, max_row=rtype_end))
        ch.set_categories(Reference(ws, min_col=2, min_row=rtype_start, max_row=rtype_end))
        ch.series[0].title = SeriesLabel(v="Parties")
        ch.series[0].graphicalProperties.solidFill = C["blue"]
        ch.legend = None
        ws.add_chart(ch, f"E{rtype_hdr}")

    if any(data["sig_counts"].values()):
        pie = PieChart()
        pie.style = 10
        pie.title = "Significance"
        pie.width = 9
        pie.height = 8
        pie.add_data(Reference(ws, min_col=9, min_row=sig_start, max_row=sig_end))
        pie.set_categories(Reference(ws, min_col=8, min_row=sig_start, max_row=sig_end))
        for idx, color in enumerate(sig_colors):
            pt = DataPoint(idx=idx)
            pt.graphicalProperties.solidFill = color
            pie.series[0].dPt.append(pt)
        ws.add_chart(pie, f"J{rtype_hdr}")

    row = max(row, rtype_hdr + int(max(7, len(sorted_rt) * 1.4) / 15 * 20) + 2)
    _row_height(ws, row, 10)
    row += 1

    # --- 5. Findings trend by review year ---------------------------------
    if len(data["year_trend"]) > 1:
        _section_header(row, 2, 11, "\u25fc  FINDINGS TREND BY REVIEW YEAR")
        row += 1
        _row_height(ws, row, 5)
        row += 1
        trend_hdr = row
        for col, (txt, align) in zip([2, 3, 4], [("Review Year", "center"), ("Findings", "center"), ("% of Total", "center")]):
            _set(ws.cell(row, col), value=txt, bold=True, fg=C["navy"],
                 font_color=C["white"], h_align=align, size=10)
        _row_height(ws, row, 18)
        row += 1
        trend_start = row
        for i, (yr, cnt) in enumerate(data["year_trend"]):
            bg  = C["grey"] if i % 2 == 0 else C["white"]
            pct = f"{cnt / total_fn * 100:.1f}%"
            _set(ws.cell(row, 2), value=yr,  fg=bg, h_align="center", size=10)
            _set(ws.cell(row, 3), value=cnt, fg=bg, h_align="center", size=10, bold=True)
            _set(ws.cell(row, 4), value=pct, fg=bg, h_align="center", size=10)
            _row_height(ws, row, 17)
            row += 1
        trend_end = row - 1

        ch2 = BarChart()
        ch2.type = "col"
        ch2.style = 10
        ch2.title = "Findings by Review Year"
        ch2.y_axis.title = "Findings"
        ch2.width = 14
        ch2.height = 9
        ch2.add_data(Reference(ws, min_col=3, min_row=trend_start, max_row=trend_end))
        ch2.set_categories(Reference(ws, min_col=2, min_row=trend_start, max_row=trend_end))
        ch2.series[0].graphicalProperties.solidFill = C["navy"]
        ch2.legend = None
        ws.add_chart(ch2, f"F{trend_hdr}")
        row = max(row, trend_hdr + 15)
        _row_height(ws, row, 10)
        row += 1

    # --- 6. Parties with / without findings by paragraph group ------------
    if data["group_submission_table"]:
        _section_header(row, 2, 11, "\u25fc  PARTIES WITH / WITHOUT FINDINGS BY PARA GROUP", color=C["green"])
        row += 1
        _row_height(ws, row, 5)
        row += 1
        sub_hdr = row
        for col, (txt, align) in zip([2, 3, 4, 5, 6], [
            ("Para Group", "left"), ("With Findings", "center"),
            ("Without", "center"), ("Total", "center"), ("% With", "center"),
        ]):
            _set(ws.cell(row, col), value=txt, bold=True, fg=C["green"],
                 font_color=C["white"], h_align=align, size=10)
        _row_height(ws, row, 18)
        row += 1
        sub_start = row
        for i, (pg, with_s, without, total) in enumerate(data["group_submission_table"]):
            bg  = C["grey"] if i % 2 == 0 else C["white"]
            pct = f"{with_s / total * 100:.0f}%" if total else "-"
            _set(ws.cell(row, 2), value=pg,      fg=bg, h_align="left",   size=10)
            _set(ws.cell(row, 3), value=with_s,  fg=bg, h_align="center", size=10, bold=True, font_color=C["green"])
            _set(ws.cell(row, 4), value=without, fg=bg, h_align="center", size=10, font_color=C["red"])
            _set(ws.cell(row, 5), value=total,   fg=bg, h_align="center", size=10)
            _set(ws.cell(row, 6), value=pct,     fg=bg, h_align="center", size=10)
            _row_height(ws, row, 17)
            row += 1
        sub_end = row - 1

        ch3 = BarChart()
        ch3.type = "bar"
        ch3.grouping = "stacked"
        ch3.style = 10
        ch3.title = "With vs Without Findings by Group"
        ch3.width = 13
        ch3.height = max(7, len(data["group_submission_table"]) * 1.4)
        for ci, (lbl, color) in enumerate([("With", C["green"]), ("Without", C["red_light"])], start=3):
            ch3.add_data(Reference(ws, min_col=ci, min_row=sub_start, max_row=sub_end))
            ch3.series[ci - 3].title = SeriesLabel(v=lbl)
            ch3.series[ci - 3].graphicalProperties.solidFill = color
        ch3.set_categories(Reference(ws, min_col=2, min_row=sub_start, max_row=sub_end))
        ws.add_chart(ch3, f"H{sub_hdr}")
        row = max(row, sub_hdr + int(max(7, len(data["group_submission_table"]) * 1.4) / 15 * 20) + 2)
        _row_height(ws, row, 10)
        row += 1

    # --- 7. Inconsistencies by paragraph group ----------------------------
    _section_header(row, 2, 11, "\u25fc  INCONSISTENCIES BY PARAGRAPH GROUP")
    row += 1
    _row_height(ws, row, 5)
    row += 1
    para_hdr = row
    for col, (txt, align) in zip([2, 3, 4], [("Para Group", "left"), ("Count", "center"), ("% of Total", "center")]):
        _set(ws.cell(row, col), value=txt, bold=True, fg=C["navy"],
             font_color=C["white"], h_align=align, size=10)
    _row_height(ws, row, 18)
    row += 1
    para_start = row
    for i, (pg, cnt) in enumerate(data["para_table"]):
        bg  = C["grey"] if i % 2 == 0 else C["white"]
        pct = f"{cnt / total_fn * 100:.1f}%"
        _set(ws.cell(row, 2), value=pg,  fg=bg, h_align="left",   size=10)
        _set(ws.cell(row, 3), value=cnt, fg=bg, h_align="center", size=10, bold=True)
        _set(ws.cell(row, 4), value=pct, fg=bg, h_align="center", size=10, font_color=C["text_mid"])
        _row_height(ws, row, 17)
        row += 1
    para_end = row - 1
    if data["para_table"]:
        ch4 = BarChart()
        ch4.type = "bar"
        ch4.style = 10
        ch4.title = "Findings by Para Group"
        ch4.y_axis.title = "Para Group"
        ch4.x_axis.title = "Count"
        ch4.width = 13
        ch4.height = max(7, len(data["para_table"]) * 1.2)
        ch4.add_data(Reference(ws, min_col=3, min_row=para_start, max_row=para_end))
        ch4.set_categories(Reference(ws, min_col=2, min_row=para_start, max_row=para_end))
        ch4.series[0].graphicalProperties.solidFill = C["blue"]
        ch4.legend = None
        ws.add_chart(ch4, f"F{para_hdr}")
        row = max(row, para_hdr + int(max(7, len(data["para_table"]) * 1.2) / 15 * 20) + 2)
    _row_height(ws, row, 10)
    row += 1

    # --- 8. Severity-scored party review outcomes -------------------------
    _section_header(row, 2, 11, "\u25fc  PARTY REVIEW OUTCOMES  (severity-scored)")
    row += 1
    _row_height(ws, row, 5)
    row += 1
    party_hdr = row
    for col, (txt, align) in zip([2, 3, 4, 5, 6], [
        ("Party", "left"), ("Findings", "center"),
        ("CB Needs", "center"), ("Total", "center"), ("Severity", "center"),
    ]):
        _set(ws.cell(row, col), value=txt, bold=True, fg=C["navy"],
             font_color=C["white"], h_align=align, size=10)
    _row_height(ws, row, 18)
    row += 1

    # Severity = findings * 2 + CB needs * 1 (a simple weighted score)
    max_sev = max((fn * 2 + cb for _, fn, cb in data["party_table"]), default=1) or 1
    party_start = row
    for i, (party, fn_cnt, cb_cnt) in enumerate(
        sorted(data["party_table"], key=lambda x: -(x[1] * 2 + x[2]))
    ):
        bg    = C["grey"] if i % 2 == 0 else C["white"]
        total = fn_cnt + cb_cnt
        sev   = fn_cnt * 2 + cb_cnt
        sev_pct = sev / max_sev
        sev_col = C["red"] if sev_pct > 0.66 else (C["gold"] if sev_pct > 0.33 else C["green"])
        sev_lbl = "\u25a0" * min(int(sev_pct * 5) + 1, 5)
        _set(ws.cell(row, 2), value=party,   fg=bg, h_align="left",   size=10)
        _set(ws.cell(row, 3), value=fn_cnt,  fg=bg, h_align="center", size=10, bold=True, font_color=C["red"])
        _set(ws.cell(row, 4), value=cb_cnt,  fg=bg, h_align="center", size=10, bold=True, font_color=C["green"])
        _set(ws.cell(row, 5), value=total,   fg=bg, h_align="center", size=10)
        _set(ws.cell(row, 6), value=sev_lbl, fg=bg, h_align="left",   size=10, font_color=sev_col)
        _row_height(ws, row, 17)
        row += 1
    party_end = row - 1

    if data["party_table"]:
        ch5 = BarChart()
        ch5.type = "col"
        ch5.style = 10
        ch5.title = "Findings & CB Needs by Party"
        ch5.y_axis.title = "Count"
        ch5.width = 14
        ch5.height = 10
        for ci, (lbl, color) in enumerate([("Findings", C["blue"]), ("CB Needs", C["green"])], start=3):
            ch5.add_data(Reference(ws, min_col=ci, min_row=party_start, max_row=party_end))
            ch5.series[ci - 3].title = SeriesLabel(v=lbl)
            ch5.series[ci - 3].graphicalProperties.solidFill = color
        ch5.set_categories(Reference(ws, min_col=2, min_row=party_start, max_row=party_end))
        ws.add_chart(ch5, f"H{party_hdr}")
    row = max(row, party_hdr + 22)
    _row_height(ws, row, 10)
    row += 1

    # --- 9. Issue type breakdown ------------------------------------------
    _section_header(row, 2, 11, "\u25fc  ISSUE TYPE BREAKDOWN")
    row += 1
    _row_height(ws, row, 5)
    row += 1
    issue_hdr = row
    for col, (txt, align) in zip([2, 3, 4], [("Issue Type", "left"), ("Count", "center"), ("Share", "center")]):
        _set(ws.cell(row, col), value=txt, bold=True, fg=C["blue"],
             font_color=C["white"], h_align=align, size=10)
    _row_height(ws, row, 18)
    row += 1
    issue_start = row
    for i, (itype, cnt) in enumerate(data["issue_table"]):
        bg  = C["grey"] if i % 2 == 0 else C["white"]
        pct = f"{cnt / total_fn * 100:.1f}%"
        _set(ws.cell(row, 2), value=itype, fg=bg, h_align="left",   size=10, wrap=True)
        _set(ws.cell(row, 3), value=cnt,   fg=bg, h_align="center", size=10, bold=True)
        _set(ws.cell(row, 4), value=pct,   fg=bg, h_align="center", size=10)
        _row_height(ws, row, 17)
        row += 1
    issue_end = row - 1

    if data["issue_table"]:
        ch6 = BarChart()
        ch6.type = "bar"
        ch6.style = 10
        ch6.title = "Issue Type Distribution"
        ch6.width = 13
        ch6.height = max(6, len(data["issue_table"]) * 1.3)
        ch6.add_data(Reference(ws, min_col=3, min_row=issue_start, max_row=issue_end))
        ch6.set_categories(Reference(ws, min_col=2, min_row=issue_start, max_row=issue_end))
        ch6.series[0].graphicalProperties.solidFill = C["gold"]
        ch6.legend = None
        ws.add_chart(ch6, f"F{issue_hdr}")
        row = max(row, issue_hdr + int(max(6, len(data["issue_table"]) * 1.3) / 15 * 20) + 2)
    _row_height(ws, row, 10)
    row += 1

    # --- 10. Top-10 requirements leaderboard ------------------------------
    _section_header(row, 2, 11, "\u25fc  TOP 10 REQUIREMENTS  (most findings)")
    row += 1
    _row_height(ws, row, 5)
    row += 1
    _col_width(ws, 4, 52)
    for col, (txt, align) in zip([2, 3, 4, 5], [
        ("Rank", "center"), ("Req ID", "center"), ("Label", "left"), ("Count", "center"),
    ]):
        _set(ws.cell(row, col), value=txt, bold=True, fg=C["navy"],
             font_color=C["white"], h_align=align, size=10)
    _row_height(ws, row, 18)
    row += 1
    for i, (rid, lbl, cnt) in enumerate(data["top10_req"]):
        bg    = C["grey"] if i % 2 == 0 else C["white"]
        medal = ["\U0001f947", "\U0001f948", "\U0001f949"][i] if i < 3 else f"#{i + 1}"
        _set(ws.cell(row, 2), value=medal, fg=bg, h_align="center", size=10)
        _set(ws.cell(row, 3), value=rid,   fg=bg, h_align="center", size=10)
        _set(ws.cell(row, 4), value=lbl,   fg=bg, h_align="left",   size=10, wrap=True)
        _set(ws.cell(row, 5), value=cnt,   fg=bg, h_align="center", size=11, bold=True,
             font_color=C["red"] if cnt >= 5 else C["text_dark"])
        _row_height(ws, row, 18)
        row += 1
    _row_height(ws, row, 10)
    row += 1

    # --- 11. Requirement coverage matrix ----------------------------------
    if data["matrix_parties"] and data["top_req_ids"]:
        _section_header(row, 2, 11,
                        "\u25fc  REQUIREMENT COVERAGE MATRIX  (Party \u00d7 Top Req IDs)",
                        color=C["blue"])
        row += 1
        _row_height(ws, row, 5)
        row += 1
        req_ids = data["top_req_ids"]
        _set(ws.cell(row, 2), value="Party", bold=True, fg=C["blue"],
             font_color=C["white"], size=10)
        for j, rid in enumerate(req_ids):
            _set(ws.cell(row, 3 + j), value=rid, bold=True, fg=C["blue"],
                 font_color=C["white"], h_align="center", size=8)
            _col_width(ws, 3 + j, 9)
        _set(ws.cell(row, 3 + len(req_ids)), value="Total", bold=True,
             fg=C["blue"], font_color=C["white"], h_align="center", size=10)
        _row_height(ws, row, 20)
        row += 1

        def _coverage_color(v):
            if v == 0:
                return C["white"]
            return C["blue_light"] if v == 1 else (C["gold_light"] if v <= 3 else C["red_light"])

        for i, party in enumerate(data["matrix_parties"]):
            bg = C["grey"] if i % 2 == 0 else C["white"]
            _set(ws.cell(row, 2), value=party, fg=bg, size=10)
            row_total = 0
            for j, rid in enumerate(req_ids):
                v = data["party_req_matrix"][party].get(rid, 0)
                row_total += v
                _set(ws.cell(row, 3 + j), value=(v if v else ""),
                     fg=_coverage_color(v), h_align="center", size=9, bold=(v > 0),
                     font_color=C["red"] if v > 3 else C["text_dark"])
            _set(ws.cell(row, 3 + len(req_ids)), value=row_total,
                 fg=bg, h_align="center", bold=True, size=10)
            _row_height(ws, row, 17)
            row += 1

        # Legend
        _row_height(ws, row, 5)
        row += 1
        for j, (clr, lbl) in enumerate([
            (C["blue_light"], "1 finding"),
            (C["gold_light"], "2-3"),
            (C["red_light"],  "4+"),
        ]):
            ws.cell(row, 2 + j * 2).fill = _fill(clr)
            _set(ws.cell(row, 3 + j * 2), value=lbl, size=8, font_color=C["text_mid"])
        _row_height(ws, row, 14)
        row += 1
        _row_height(ws, row, 10)
        row += 1

    # --- 12. CB needs heat-map --------------------------------------------
    _section_header(row, 2, 11,
                    "\u25fc  CAPACITY BUILDING NEEDS HEAT-MAP  (Party \u00d7 Para Group)")
    row += 1
    _row_height(ws, row, 5)
    row += 1
    if data["cb_heatmap"] and data["para_groups_cb"]:
        pg_list    = data["para_groups_cb"]
        parties_cb = sorted(data["cb_heatmap"].keys())
        _set(ws.cell(row, 2), value="Party", bold=True, fg=C["navy"],
             font_color=C["white"], size=10)
        for j, pg in enumerate(pg_list):
            _set(ws.cell(row, 3 + j), value=pg, bold=True, fg=C["navy"],
                 font_color=C["white"], h_align="center", size=9)
            _col_width(ws, 3 + j, 11)
        _set(ws.cell(row, 3 + len(pg_list)), value="Total", bold=True,
             fg=C["navy"], font_color=C["white"], h_align="center", size=10)
        _row_height(ws, row, 18)
        row += 1
        all_values = [v for pd in data["cb_heatmap"].values() for v in pd.values()]
        max_val = max(all_values) if all_values else 1

        def _heat_color(v, mx):
            if v == 0:
                return C["white"]
            ratio = v / mx
            return "D9EFD7" if ratio < 0.33 else ("FFE5A0" if ratio < 0.66 else C["red_light"])

        for i, party in enumerate(parties_cb):
            bg = C["grey"] if i % 2 == 0 else C["white"]
            _set(ws.cell(row, 2), value=party, fg=bg, size=10)
            row_total = 0
            for j, pg in enumerate(pg_list):
                v = data["cb_heatmap"][party].get(pg, 0)
                row_total += v
                _set(ws.cell(row, 3 + j), value=(v if v else ""),
                     fg=_heat_color(v, max_val), h_align="center", size=10, bold=(v > 0),
                     font_color=C["red"] if v >= max_val * 0.66 else C["text_dark"])
            _set(ws.cell(row, 3 + len(pg_list)), value=row_total,
                 fg=bg, h_align="center", bold=True, size=10)
            _row_height(ws, row, 17)
            row += 1
    else:
        _merge_set(ws, row, 2, row, 11,
                   value="  No CB needs data yet.", fg=C["grey"],
                   font_color=C["text_mid"], italic=True, size=10)
        _row_height(ws, row, 18)
        row += 1
    _row_height(ws, row, 10)
    row += 1

    # --- 13. Report type distribution -------------------------------------
    if data["rtype_counts"]:
        _section_header(row, 2, 11, "\u25fc  REPORT TYPE DISTRIBUTION")
        row += 1
        _row_height(ws, row, 5)
        row += 1
        for col, (txt, align) in zip([2, 3, 4], [("Report Type", "left"), ("TER Count", "center"), ("Share", "center")]):
            _set(ws.cell(row, col), value=txt, bold=True, fg=C["blue"],
                 font_color=C["white"], h_align=align, size=10)
        _row_height(ws, row, 18)
        row += 1
        total_ters = data["n_ters"] or 1
        for i, (rtype, cnt) in enumerate(sorted(data["rtype_counts"].items(), key=lambda x: -x[1])):
            bg  = C["grey"] if i % 2 == 0 else C["white"]
            pct = f"{cnt / total_ters * 100:.1f}%"
            _set(ws.cell(row, 2), value=rtype, fg=bg, size=10)
            _set(ws.cell(row, 3), value=cnt,   fg=bg, h_align="center", size=10, bold=True)
            _set(ws.cell(row, 4), value=pct,   fg=bg, h_align="center", size=10)
            _row_height(ws, row, 17)
            row += 1
        _row_height(ws, row, 10)
        row += 1

    # --- 14. Footer -------------------------------------------------------
    _row_height(ws, row, 8)
    row += 1
    for c in range(1, 13):
        ws.cell(row, c).fill = _fill(C["navy"])
    _merge_set(ws, row, 2, row, 11,
               value=f"  UNFCCC Article 6.2 TER Dashboard  \u2022  Generated {date.today().strftime('%d %B %Y')}",
               fg=C["navy"], font_color="5A7FA8", h_align="left", size=8, italic=True)
    _row_height(ws, row, 20)

    ws.freeze_panes = "B7"

    _log(f"   Dashboard built: {row} rows | "
         f"{data['n_parties']} parties | {data['n_findings']} findings | "
         f"FDR complete {data['fdr_rate']}% | Rec rate {data['rec_rate']}%")

    return ws


# ---------------------------------------------------------------------------
# Sheet utilities (used by the write pipeline)
# ---------------------------------------------------------------------------

def sheet_headers(ws) -> dict[str, int]:
    """Return {header_name: column_index} keeping only the first occurrence of each name."""
    headers = {}
    for col in range(1, ws.max_column + 1):
        name = ws.cell(1, col).value
        if name and name not in headers:
            headers[name] = col
    return headers


def true_last_row(ws) -> int:
    """Return the last row that actually contains any data (openpyxl's max_row can overshoot)."""
    last = 1
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            if cell.value is not None:
                last = cell.row
                break
    return last


def next_id_num(ws, id_col_name: str) -> int:
    """Return the next serial number to use based on the highest trailing integer found in the ID column."""
    headers = sheet_headers(ws)
    col = headers.get(id_col_name, 1)
    max_num = 0
    for row in ws.iter_rows(min_row=2, max_col=col, values_only=True):
        val = row[col - 1]
        if isinstance(val, str):
            match = re.search(r"(\d+)$", val)
            if match:
                max_num = max(max_num, int(match.group(1)))
        elif isinstance(val, (int, float)):
            max_num = max(max_num, int(val))
    return max_num + 1


def append_row_by_name(ws, headers_map: dict[str, int], values: dict[str, Any]):
    """Append a new row, matching values to columns by header name."""
    next_row = true_last_row(ws) + 1
    for col_name, col_idx in headers_map.items():
        val = values.get(col_name, "")
        ws.cell(next_row, col_idx).value = val if val != "" else None


# ---------------------------------------------------------------------------
# Reference lookups (ref_Requirements, ref_Parties)
# ---------------------------------------------------------------------------

def build_req_lookup(wb) -> dict[str, dict[str, str]]:
    """Build a {Req_ID: metadata} lookup from the ref_Requirements sheet."""
    lookup: dict[str, dict[str, str]] = {}
    if "ref_Requirements" not in wb.sheetnames:
        return lookup

    ws = wb["ref_Requirements"]
    h  = sheet_headers(ws)
    col_req_id   = h.get("Req_ID",                    h.get("A6_2_ID", 1))
    col_para     = h.get("Para_Group",                2)
    col_sub      = h.get("Sub_Para",                  3)
    col_info_lvl = h.get("Info_Level",                6)
    col_short    = h.get("Requirement_Short_Label",   h.get("Short_Label"))
    col_req_text = h.get("Requirement_Text",          9)
    col_section  = h.get("Report_Section",            10)
    col_review   = h.get("Review_Section2",           col_section)

    def _cell(row, col):
        return str(row[col - 1] if col and col <= len(row) else "") or ""

    for row in ws.iter_rows(min_row=2, values_only=True):
        req_id = row[col_req_id - 1] if col_req_id <= len(row) else None
        if not req_id:
            continue
        req_text = _cell(row, col_req_text)
        short_label = _cell(row, col_short) if col_short else _cell(row, col_review)
        if not short_label:
            short_label = req_text[:60]
        lookup[str(req_id).strip()] = {
            "Para_Group":       _cell(row, col_para),
            "Sub_Para":         _cell(row, col_sub),
            "Info_Level":       _cell(row, col_info_lvl),
            "Short_Label":      short_label,
            "Requirement_Text": req_text,
            "Report_Section":   _cell(row, col_section),
        }
    return lookup


def build_party_lookup(wb) -> dict[str, str]:
    """Build a {party_name_lower: party_code} lookup from the ref_Parties sheet."""
    lookup: dict[str, str] = {}
    if "ref_Parties" not in wb.sheetnames:
        return lookup

    ws = wb["ref_Parties"]
    h  = sheet_headers(ws)
    col_code = h.get("Party_Code", 1)
    col_name = h.get("Party_Name", 2)

    for row in ws.iter_rows(min_row=2, values_only=True):
        code = row[col_code - 1] if col_code <= len(row) else None
        name = row[col_name - 1] if col_name <= len(row) else None
        if code and name:
            lookup[str(name).strip().lower()] = str(code).strip()
    return lookup


def lookup_req(req_id: str, req_lookup: dict) -> dict:
    """Look up a requirement, retrying with whitespace stripped if the exact key is not found."""
    info = req_lookup.get(req_id, {})
    if info:
        return info
    normalised = req_id.replace(" ", "")
    for key, value in req_lookup.items():
        if key.replace(" ", "") == normalised:
            return value
    return {}


# ---------------------------------------------------------------------------
# Word document parsing
# ---------------------------------------------------------------------------

def extract_doc_symbol(doc) -> str:
    """Pull the FCCC document symbol out of the first couple of tables."""
    for table in doc.tables[:2]:
        for row in table.rows:
            for cell in row.cells:
                match = re.search(r"FCCC/[A-Z0-9][A-Z0-9/\.\-]+", cell.text)
                if match:
                    return match.group(0).strip()
    return ""


def extract_party_info(doc, doc_symbol: str, party_lookup: dict) -> tuple[str, str]:
    """Return (party_name, party_code) inferred from the document text and ref_Parties."""
    full_text = "\n".join(p.text for p in doc.paragraphs)
    match = re.search(
        r"(?:initial report|updated initial report|regular information)"
        r".*?of\s+([A-Z][a-zA-Z\s\-]+?)(?:\n|,|\.|$)",
        full_text, re.IGNORECASE,
    )
    party_name = match.group(1).strip() if match else "Unknown"
    party_code = party_lookup.get(party_name.lower(), "")

    if not party_code:
        # Fall back to the 3-letter code embedded in the document symbol
        code_match = re.search(r"/([A-Z]{3})(?:/|Add|\.)", doc_symbol)
        party_code = code_match.group(1) if code_match else party_name[:3].upper()

    return party_name, party_code


def extract_report_type(doc) -> str:
    """Identify the report type from phrases in the document body."""
    text = "\n".join(p.text for p in doc.paragraphs).lower()
    if "updated initial report" in text:
        return "UIR-Updated Initial Report"
    if "regular information" in text:
        return "RI-Regular Information"
    if "annual information" in text:
        return "AI-Annual Information"
    if "biennial transparency" in text:
        return "BTR-Biennial Transparency Report"
    return "IR-Initial Report"


def extract_review_info(doc) -> tuple[str, str]:
    """Return (review_period, review_year), defaulting year to the current year."""
    full_text = "\n".join(p.text for p in doc.paragraphs)
    review_period = ""
    review_year = str(date.today().year)

    match = re.search(r"took place from (.+?) in (.+?)[\.\n]", full_text)
    if match:
        review_period = match.group(1).strip()
        year_match = re.search(r"20\d{2}", review_period)
        if year_match:
            review_year = year_match.group(0)

    return review_period, review_year


def extract_ca_from_paragraphs(doc) -> tuple[str, str]:
    """Return (CA ID, CA name) if the cooperative approach is declared inline."""
    for paragraph in doc.paragraphs:
        match = re.search(
            r"cooperative approach\s+(CA\d+),\s*[\u201c\u201d\"](.+?)[\u201c\u201d\"]",
            paragraph.text, re.IGNORECASE,
        )
        if match:
            return match.group(1).strip(), match.group(2).strip()
    return "", ""


# ---------------------------------------------------------------------------
# FDR (Finding / During / Recommendation) splitter
# ---------------------------------------------------------------------------

AI_SYSTEM_PROMPT = """You are a data extraction assistant for UNFCCC Article 6.2 Technical Expert Review Reports (TERRs).

Your task: parse the Description column of a TERR findings table into structured FDR blocks.

Each cell may contain one or more FDR blocks. Patterns you may encounter:
- Simple FDR: Finding -> During -> Recommendation (most common)
- Numbered FDR-FDR: (1) Finding -> During -> Recommendation (2) Finding -> During -> Recommendation
- F-DR-DR: one Finding followed by two During-Recommendation pairs
- Sig+FDR: significance determination paragraph appears BEFORE the During section
- FDR+Sig-at-end: significance paragraph appears AFTER the Recommendation
- Missing During or Recommendation sections (partial blocks)

CRITICAL RULES:
- "Finding" = what the Party reported + what was found insufficient/inconsistent. INCLUDES any significance determination paragraph ("The Article 6 TERT has determined this inconsistency as significant..."). Everything before "During the review".
- "During" = starts with "During the review" - what happened during the review discussion. Ends before the real Recommendation.
- "Recommendation" = ALL paragraphs starting with "The Article 6 TERT acknowledges...", "The Article 6 TERT recommends...", OR "The Article 6 TERT also recommends...". Multiple recommendation/acknowledgement sentences for the SAME FDR block must be COMBINED into one rec_text string separated by newlines - do NOT create separate blocks for each sentence.
- Do NOT include "The Article 6 TERT has determined this inconsistency as significant..." in the Recommendation - that belongs in the Finding.
- Numbered blocks like (1)...(2)... are SEPARATE FDR entries -> return as separate objects
- issue_type: one of "Cross-party inconsistency", "Quantified inconsistency", "Cross-report inconsistency", "Qualitative inconsistency", or "" if not mentioned
- significance: "Significant" if text says "determined this inconsistency as significant", "Minor" if minor, "" if not mentioned

Return ONLY valid JSON - no markdown, no explanation:
{
  "blocks": [
    {
      "finding_text": "...",
      "during_text": "...",
      "rec_text": "...",
      "significance": "Significant|Minor|",
      "issue_type": "..."
    }
  ]
}"""


def _extract_sig_and_issue(text: str) -> tuple[str, str]:
    """Pull the significance label and issue type out of a block of text."""
    sig_match = re.search(
        r"(?:inconsistency (?:is|as)|determined (?:this |that this )?inconsistency as)\s*(significant|minor)",
        text, re.IGNORECASE,
    )
    issue_match = re.search(
        r"(Cross-party inconsistency|Quantified inconsistency"
        r"|Cross-report inconsistency|Qualitative inconsistency)",
        text, re.IGNORECASE,
    )
    return (
        sig_match.group(1).capitalize() if sig_match else "",
        issue_match.group(1) if issue_match else "",
    )


def _classify_paragraph(paragraph: str) -> str:
    """Classify a paragraph as F (Finding), D (During), R (Recommendation), SIG, or EMPTY."""
    text = paragraph.strip()
    if not text:
        return "EMPTY"

    # Recommendation: acknowledges / recommends / notes / encourages / urges
    if re.match(
        r"The Article 6 TERT\b.{0,20}\b(acknowledges|recommends|notes that|encourages|urges)",
        text, re.I,
    ):
        return "R"

    # Significance determination paragraph
    if re.match(r"The Article 6 TERT (?:has determined|considers that)", text, re.I):
        return "SIG"

    # During-the-review discussion
    if re.match(r"During the review", text, re.I):
        return "D"

    return "F"


def _regex_split(text: str) -> list[dict]:
    """
    Paragraph-level FDR classifier used when AI parsing is unavailable or unnecessary.

    Groups consecutive paragraphs: F/SIG -> D -> R(s) form one block.
    Multiple consecutive R paragraphs (e.g. 'acknowledges' + 'also recommends')
    are concatenated into a single Recommendation_Text; a new block is only
    started when an F/D paragraph appears after at least one R.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    sig, issue = _extract_sig_and_issue(text)

    if not lines:
        return [{
            "finding_text": "", "during_text": "", "rec_text": "",
            "significance": sig, "issue_type": issue,
        }]

    blocks: list[dict] = []
    cur_f: list[str] = []
    cur_d: list[str] = []
    cur_r: list[str] = []
    state = "F"

    def flush():
        nonlocal cur_f, cur_d, cur_r
        if cur_f or cur_d or cur_r:
            blocks.append({
                "finding_text": "\n".join(cur_f).strip(),
                "during_text":  "\n".join(cur_d).strip(),
                "rec_text":     "\n".join(cur_r).strip(),
            })
        cur_f, cur_d, cur_r = [], [], []

    for line in lines:
        tag = _classify_paragraph(line)
        if tag == "EMPTY":
            continue

        if tag in ("F", "SIG"):
            if state == "R" and cur_r:
                flush()
                state = "F"
            cur_f.append(line)
            state = "F"
        elif tag == "D":
            if state == "R" and cur_r:
                flush()
                state = "F"
            cur_d.append(line)
            state = "D"
        elif tag == "R":
            # Keep appending into the current Recommendation, never split on consecutive R lines
            cur_r.append(line)
            state = "R"

    flush()

    if not blocks:
        return [{
            "finding_text": text.strip(), "during_text": "", "rec_text": "",
            "significance": sig, "issue_type": issue,
        }]

    # Only attach sig/issue_type to the first block; the subsequent blocks inherit nothing
    return [
        {**b, "significance": sig if i == 0 else "", "issue_type": issue if i == 0 else ""}
        for i, b in enumerate(blocks)
    ]


def _needs_ai(text: str) -> bool:
    """Return True if the FDR text has structure complex enough to warrant AI parsing."""
    # Numbered blocks like (2), (3)...
    if re.search(r"^\s*\(2\)", text, re.MULTILINE):
        return True
    # Multiple "During the review" markers
    if len(re.findall(r"During the review", text, re.IGNORECASE)) > 1:
        return True
    # Significance paragraph appearing before "During" (the sig+FDR pattern)
    sig_pos = text.lower().find("the article 6 tert has determined this inconsistency")
    dur_pos = text.lower().find("during the review")
    if sig_pos != -1 and dur_pos != -1 and sig_pos < dur_pos:
        return True
    return False


def split_description_ai(text: str, log=None) -> list[dict]:
    """
    Split a Description cell into FDR blocks.

    Uses regex for simple cases (no API call needed). For complex cases it
    calls the Anthropic API if ANTHROPIC_API_KEY is set; otherwise falls
    back to regex.
    """
    if not text.strip():
        return [{
            "finding_text": "", "during_text": "", "rec_text": "",
            "significance": "", "issue_type": "",
        }]

    if not _needs_ai(text):
        return _regex_split(text)

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        if log:
            log("         ANTHROPIC_API_KEY not set - using regex fallback")
        return _regex_split(text)

    try:
        payload = json.dumps({
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 4000,
            "system": AI_SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": f"Parse this TERR description:\n\n{text}"}],
        }).encode()
        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read())

        raw = data["content"][0]["text"].strip()
        # Strip any markdown fences the model might emit
        raw = re.sub(r"^```json\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

        parsed = json.loads(raw)
        blocks = parsed.get("blocks", [])
        if blocks and log:
            log(f"         AI split -> {len(blocks)} block(s)")
        return blocks if blocks else _regex_split(text)

    except Exception as exc:
        if log:
            log(f"         AI parse failed ({exc}), using regex fallback")
        return _regex_split(text)


# ---------------------------------------------------------------------------
# Word table detection and row extraction
# ---------------------------------------------------------------------------

def detect_table_format(table) -> str:
    """Identify which TERR table layout we are looking at."""
    if len(table.rows) < 2:
        return "unknown"

    headers = [c.text.strip().lower() for c in table.rows[0].cells]
    h0 = headers[0]
    h1 = headers[1] if len(headers) > 1 else ""
    h2 = headers[2] if len(headers) > 2 else ""

    # Format A - first column is 'id#'
    if h0 in ("id#", "id #", "id"):
        if "area in which capacity" in h2 or "capacity-building" in h2:
            return "format_a_cb"
        if "description" in h2 or "element of information" in h1:
            return "format_a_findings"

    # Format B - first column is a requirement sentence
    if any(keyword in h0 for keyword in ("requirement", "demonstrate", "provide", "describe")):
        # B2 = 5+ columns with Significant + Persistent columns
        if len(headers) >= 4 and any("significant" in h for h in headers[3:]):
            return "format_b2"
        return "format_b"

    return "unknown"


def _norm_yes_no(val: Any) -> str:
    """Normalise variants of Yes/No/NA into canonical strings."""
    if not val:
        return ""
    v = str(val).strip().lower()
    if v in ("yes", "y", "true"):
        return "Yes"
    if v in ("no", "n", "false"):
        return "No"
    if v in ("na", "n/a", "not applicable"):
        return "NA"
    return str(val).strip()


def _make_format_b_row(raw_req_id, req_info, ca_id_global, ca_name_global,
                       significance="", persistent=""):
    """Build a row dict for a Format B / B2 requirement."""
    info_level = req_info.get("Info_Level", "")
    return {
        "req_id":           raw_req_id,
        "para_group":       req_info.get("Para_Group", ""),
        "sub_para":         req_info.get("Sub_Para", ""),
        "paragraph_title":  req_info.get("Short_Label", ""),
        "requirement_text": req_info.get("Requirement_Text", ""),
        "info_level":       info_level,
        "ca_id":            ca_id_global   if info_level == "CA-specific" else "",
        "ca_name":          ca_name_global if info_level == "CA-specific" else "",
        "finding_text":  "", "during_text": "", "rec_text": "",
        "significance":  significance,
        "issue_type":    "",
        "persistent":    persistent,
        "doc_format":    "B",
    }


def process_format_b(table, req_lookup, ca_id_global, ca_name_global) -> list[dict]:
    """Extract rows from a 3-4 column Format B table: Requirement | Element | ID#."""
    rows: list[dict] = []
    req_id_pattern = re.compile(r"^18\.[A-Z0-9]")

    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        raw_req_ids = [cells[i] for i in range(2, len(cells)) if req_id_pattern.match(cells[i])]
        for raw_req_id in raw_req_ids:
            req_info = lookup_req(raw_req_id, req_lookup)
            rows.append(_make_format_b_row(raw_req_id, req_info, ca_id_global, ca_name_global))
    return rows


def process_format_b2(table, req_lookup, ca_id_global, ca_name_global) -> list[dict]:
    """Extract rows from a 5+ column Format B2 table with Significant/Persistent columns."""
    rows: list[dict] = []
    req_id_pattern = re.compile(r"^18\.[A-Z0-9]")

    headers = [c.text.strip().lower() for c in table.rows[0].cells]
    sig_col  = next((i for i, h in enumerate(headers) if "significant" in h), None)
    pers_col = next((i for i, h in enumerate(headers) if "persistent"  in h), None)

    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue

        raw_req_id = cells[2].strip() if len(cells) > 2 else ""
        if not req_id_pattern.match(raw_req_id):
            continue

        sig_val  = _norm_yes_no(cells[sig_col])  if sig_col  is not None and sig_col  < len(cells) else ""
        pers_val = _norm_yes_no(cells[pers_col]) if pers_col is not None and pers_col < len(cells) else ""

        req_info = lookup_req(raw_req_id, req_lookup)
        rows.append(_make_format_b_row(
            raw_req_id, req_info, ca_id_global, ca_name_global,
            significance=sig_val, persistent=pers_val,
        ))
    return rows


def process_format_a_findings(table, req_lookup, ca_id_global, ca_name_global, log=None) -> list[dict]:
    """Extract findings from a Format A table (ID# | Element | Description)."""
    rows: list[dict] = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2 or not cells[0].strip():
            continue

        raw_req_id = cells[0].strip()
        desc       = cells[2].strip() if len(cells) > 2 else ""
        req_info   = lookup_req(raw_req_id, req_lookup)

        info_level = req_info.get("Info_Level", "")
        ca_id   = ca_id_global   if info_level == "CA-specific" else ""
        ca_name = ca_name_global if info_level == "CA-specific" else ""

        # The description may pack one or more FDR blocks
        for block in split_description_ai(desc, log=log):
            rows.append({
                "req_id":           raw_req_id,
                "para_group":       req_info.get("Para_Group", ""),
                "sub_para":         req_info.get("Sub_Para", ""),
                "paragraph_title":  req_info.get("Short_Label", ""),
                "requirement_text": req_info.get("Requirement_Text", ""),
                "info_level":       info_level,
                "ca_id":            ca_id,
                "ca_name":          ca_name,
                "finding_text":     block.get("finding_text", ""),
                "during_text":      block.get("during_text", ""),
                "rec_text":         block.get("rec_text", ""),
                "significance":     block.get("significance", ""),
                "issue_type":       block.get("issue_type", ""),
            })
    return rows


def process_format_a_cb(table, req_lookup, ca_id_global, ca_name_global) -> list[dict]:
    """Extract CB needs from a Format A capacity-building table."""
    rows: list[dict] = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2 or not cells[0].strip():
            continue

        raw_req_id = cells[0].strip()
        desc       = cells[2].strip() if len(cells) > 2 else ""
        req_info   = lookup_req(raw_req_id, req_lookup)

        rows.append({
            "req_id":           raw_req_id,
            "para_group":       req_info.get("Para_Group", ""),
            "sub_para":         req_info.get("Sub_Para", ""),
            "paragraph_title":  req_info.get("Short_Label", ""),
            "requirement_text": req_info.get("Requirement_Text", ""),
            "info_level":       req_info.get("Info_Level", ""),
            "ca_id":            ca_id_global   if req_info.get("Info_Level") == "CA-specific" else "",
            "ca_name":          ca_name_global if req_info.get("Info_Level") == "CA-specific" else "",
            "description":      desc,
        })
    return rows


def process_word_document(docx_path: str, req_lookup: dict, party_lookup: dict, log):
    """Parse a single TERR Word document, returning (metadata, findings, cb_needs)."""
    log(f"   Reading : {os.path.basename(docx_path)}")
    doc = Document(docx_path)

    doc_symbol                   = extract_doc_symbol(doc)
    party_name, party_code       = extract_party_info(doc, doc_symbol, party_lookup)
    report_type                  = extract_report_type(doc)
    review_period, review_year   = extract_review_info(doc)
    ca_id_global, ca_name_global = extract_ca_from_paragraphs(doc)

    log(f"   Party   : {party_name} ({party_code})")
    log(f"   Symbol  : {doc_symbol or 'not detected'}")
    log(f"   Type    : {report_type}  |  Year: {review_year}")
    if ca_id_global:
        log(f"   CA      : {ca_id_global} - {ca_name_global[:50]}...")

    metadata = {
        "doc_symbol":    doc_symbol,
        "party_name":    party_name,
        "party_code":    party_code,
        "report_type":   report_type,
        "review_year":   review_year,
        "review_period": review_period,
        "ter_format":    "DR-Desk review",
        "info_type":     "Cooperative Approach",
        "language":      "English",
    }

    findings: list[dict] = []
    cb_needs: list[dict] = []
    for table in doc.tables:
        fmt = detect_table_format(table)
        if fmt == "format_a_findings":
            findings.extend(process_format_a_findings(table, req_lookup, ca_id_global, ca_name_global, log=log))
        elif fmt == "format_a_cb":
            cb_needs.extend(process_format_a_cb(table, req_lookup, ca_id_global, ca_name_global))
        elif fmt == "format_b":
            findings.extend(process_format_b(table, req_lookup, ca_id_global, ca_name_global))
        elif fmt == "format_b2":
            findings.extend(process_format_b2(table, req_lookup, ca_id_global, ca_name_global))

    log(f"   Extracted: {len(findings)} findings, {len(cb_needs)} CB needs")
    return metadata, findings, cb_needs


# ---------------------------------------------------------------------------
# Writers for the three data sheets
# ---------------------------------------------------------------------------

def write_ter_status(ws, metadata, num_findings, num_cb, serial, log) -> str:
    """Append a TER status row and return the new TER_Status_ID."""
    headers = sheet_headers(ws)
    ter_id  = f"TR-{metadata['party_code']}-{str(serial).zfill(4)}"
    append_row_by_name(ws, headers, {
        "TER_Status_ID":        ter_id,
        "Serial_Num":           serial,
        "Party_Code":           metadata["party_code"],
        "Party_Name":           metadata["party_name"],
        "Report_Type":          metadata["report_type"],
        "Language":             metadata["language"],
        "Review_Week_Start":    metadata["review_period"],
        "TERR_Document_Symbol": metadata["doc_symbol"],
        "Num_Inconsistencies":  num_findings,
        "Num_CB_Needs":         num_cb,
    })
    log(f"   tbl_TER_Status  -> {ter_id}  ({num_findings} findings, {num_cb} CB needs)")
    return ter_id


def write_findings(ws, metadata, findings, ter_id, fn_start, log) -> dict[str, str]:
    """Append finding rows and return a {req_id: finding_id} map for CB cross-referencing."""
    # Deduplicate header names: keep only the first occurrence, drop ghost duplicates
    raw_headers = sheet_headers(ws)
    headers: dict[str, int] = {}
    for name, col_idx in raw_headers.items():
        if name not in headers:
            headers[name] = col_idx

    fn_id_map: dict[str, str] = {}
    for i, row in enumerate(findings):
        fn_id = f"FN-{metadata['party_code']}-{str(fn_start + i).zfill(4)}"
        fn_id_map[row["req_id"]] = fn_id

        # Map the internal significance value to the Excel convention: Yes/No/NA
        sig_raw = row.get("significance", "")
        if sig_raw == "Significant":
            sig_out = "Yes"
        elif sig_raw == "Minor":
            sig_out = "No"
        elif sig_raw in ("NA", "N/A"):
            sig_out = "NA"
        else:
            sig_out = sig_raw or ""

        pers_out = row.get("persistent", "")

        append_row_by_name(ws, headers, {
            "Finding_ID":                fn_id,
            "TER_Status_ID":             ter_id,
            "Party_Code":                metadata["party_code"],
            "Party_Name":                metadata["party_name"],
            "Report_Type":               metadata["report_type"],
            "TER_Format":                metadata["ter_format"],
            "Review_Year":               metadata["review_year"],
            "Information_Type":          metadata["info_type"],
            "Info_Level":                row["info_level"],
            "Cooperative_Approach_ID":   row["ca_id"],
            "Cooperative_Approach_Name": row["ca_name"],
            "Requirement_ID":            row["req_id"],
            "Paragraph_Title":           row["paragraph_title"],
            "Paragraph_Group":           row["para_group"],
            "Sub_Para_Group":            row["sub_para"],
            "Requirement_Text":          row["requirement_text"],
            "Issue_Type":                row["issue_type"],
            "Significance":              sig_out,
            "Persistent_Issue":          pers_out,
            "Recommendation_Flag":       "Yes" if row["rec_text"] else "",
            "Finding_Text":              row["finding_text"],
            "During_Text":               row["during_text"],
            "Recommendation_Text":       row["rec_text"],
            "Document_Symbol":           metadata["doc_symbol"],
        })

        sig_tag  = f" sig={sig_out}"  if sig_out  else ""
        pers_tag = f" pers={pers_out}" if pers_out else ""
        fmt_tag  = " [Fmt B]" if row.get("doc_format") == "B" else ""
        log(f"      {fn_id}  [{row['req_id']}]  {row['paragraph_title'][:28]}{sig_tag}{pers_tag}{fmt_tag}")

    return fn_id_map


def write_cb_needs(ws, metadata, cb_needs, fn_id_map, cb_start, log):
    """Append capacity-building need rows."""
    headers = sheet_headers(ws)
    for i, row in enumerate(cb_needs):
        cb_id     = f"CB-{metadata['party_code']}-{str(cb_start + i).zfill(4)}"
        source_fn = fn_id_map.get(row["req_id"], "")

        append_row_by_name(ws, headers, {
            "CB_ID":                     cb_id,
            "Source_Finding_ID":         source_fn,
            "Party_Code":                metadata["party_code"],
            "Party_Name":                metadata["party_name"],
            "Report_Type":               metadata["report_type"],
            "TER_Format":                metadata["ter_format"],
            "Review_Year":               metadata["review_year"],
            "Information_Type":          metadata["info_type"],
            "Cooperative_Approach_ID":   row["ca_id"],
            "Cooperative_Approach_Name": row["ca_name"],
            "Requirement_ID":            row["req_id"],
            "Paragraph_Group":           row["para_group"],
            "Paragraph_Title":           row["paragraph_title"],
            "Requirement_Text":          row["requirement_text"],
            "CB_Need_Flag":              "Yes",
            "CB_Need_Description":       row["description"],
        })
        log(f"      {cb_id}  [{row['req_id']}]  {row['paragraph_title'][:35]}")


# ---------------------------------------------------------------------------
# Orchestrators
# ---------------------------------------------------------------------------

def sync_all(docx_paths: list[str], xlsx_path: str, log):
    """Process every Word document, write to the three tables, and refresh the dashboard."""
    log("Opening Excel database...")
    wb = load_workbook(xlsx_path)

    req_lookup   = build_req_lookup(wb)
    party_lookup = build_party_lookup(wb)
    log(f"   Reference loaded: {len(req_lookup)} requirements, {len(party_lookup)} parties")

    ws_ter = wb["tbl_TER_Status"]
    ws_fn  = wb["tbl_Metadata_Findings"]
    ws_cb  = wb["tbl_CB_Needs"]

    total_findings = 0
    total_cb = 0

    for idx, docx_path in enumerate(docx_paths):
        log(f"\n-- Document {idx + 1} of {len(docx_paths)} " + "-" * 33)
        try:
            metadata, findings, cb_needs = process_word_document(
                docx_path, req_lookup, party_lookup, log,
            )
        except Exception as exc:
            log(f"   Failed to read: {exc}")
            continue

        ter_serial = next_id_num(ws_ter, "Serial_Num")
        ter_id     = write_ter_status(ws_ter, metadata, len(findings), len(cb_needs), ter_serial, log)

        fn_start = next_id_num(ws_fn, "Finding_ID")
        log(f"\n   Writing {len(findings)} findings -> tbl_Metadata_Findings")
        fn_id_map = write_findings(ws_fn, metadata, findings, ter_id, fn_start, log)

        cb_start = next_id_num(ws_cb, "CB_ID")
        log(f"\n   Writing {len(cb_needs)} CB needs -> tbl_CB_Needs")
        write_cb_needs(ws_cb, metadata, cb_needs, fn_id_map, cb_start, log)

        total_findings += len(findings)
        total_cb       += len(cb_needs)

    build_dashboard(wb, log)
    wb.save(xlsx_path)

    log("\n" + "=" * 55)
    log("Done.")
    log(f"    Documents processed : {len(docx_paths)}")
    log(f"    TER Status rows     : {len(docx_paths)}")
    log(f"    Findings added      : {total_findings}")
    log(f"    CB Needs added      : {total_cb}")
    log(f"    Dashboard           : refreshed")
    log(f"    Saved               : {os.path.basename(xlsx_path)}")
    log("\n  Upload the updated Excel back to SharePoint when ready.")


def refresh_dashboard_only(xlsx_path: str, log):
    """Rebuild only the Dashboard sheet on an existing workbook - no Word documents needed."""
    log("Opening Excel database (dashboard-only refresh)...")
    wb = load_workbook(xlsx_path)
    build_dashboard(wb, log)
    wb.save(xlsx_path)
    log(f"Dashboard refreshed and saved -> {os.path.basename(xlsx_path)}")


# ---------------------------------------------------------------------------
# Tkinter GUI
# ---------------------------------------------------------------------------

class App(tk.Tk):
    """Simple Tk front-end: pick Word files + Excel database, run sync or dashboard-only refresh."""

    BG_DARK    = "#0a1628"
    BG_MID     = "#1e3a5f"
    BG_PANEL   = "#071020"
    FG_GOLD    = "#e8d5a3"
    FG_MUTED   = "#5a7fa8"
    FG_TEXT    = "#a8c8e8"
    ACCENT     = "#c9a84c"

    def __init__(self):
        super().__init__()
        self.title("UNFCCC A6.2 - Sync Tool")
        self.geometry("700x780")
        self.resizable(False, False)
        self.configure(bg=self.BG_DARK)

        self.docx_paths: list[str] = []
        self.xlsx_path: str | None = None

        self._build_header()
        self._build_docx_section()
        self._build_xlsx_section()
        self._build_action_buttons()
        self._build_log_panel()

    # -- Layout -----------------------------------------------------------

    def _build_header(self):
        hdr = tk.Frame(self, bg=self.BG_DARK)
        hdr.pack(fill="x", padx=30, pady=(24, 0))
        tk.Label(hdr, text="UN  Article 6.2  Sync Tool",
                 font=("Georgia", 16, "bold"), bg=self.BG_DARK, fg=self.FG_GOLD
                 ).pack(anchor="w")
        tk.Label(hdr,
                 text="Format A + Format B  |  AI-assisted FDR split  |  Auto Dashboard",
                 font=("Georgia", 9), bg=self.BG_DARK, fg=self.FG_MUTED
                 ).pack(anchor="w", pady=(3, 0))
        tk.Frame(self, bg=self.BG_MID, height=1).pack(fill="x", padx=30, pady=14)

    def _build_docx_section(self):
        section = tk.Frame(self, bg=self.BG_DARK)
        section.pack(fill="x", padx=30)

        tk.Label(section, text="1.  Word Reports (.docx) - select one or many",
                 font=("Georgia", 10, "bold"), bg=self.BG_DARK, fg=self.FG_GOLD
                 ).pack(anchor="w")
        tk.Label(section, text="    Accepts both Addendum and Main TERR formats",
                 font=("Georgia", 8), bg=self.BG_DARK, fg=self.FG_MUTED
                 ).pack(anchor="w")

        btn_row = tk.Frame(section, bg=self.BG_DARK)
        btn_row.pack(fill="x", pady=(6, 4))
        tk.Button(btn_row, text="+ Add Word Files", font=("Georgia", 9),
                  bg=self.BG_MID, fg=self.FG_GOLD, activebackground="#2a4f7f",
                  relief="flat", cursor="hand2", padx=14, pady=4,
                  command=self._add_docx).pack(side="left")
        tk.Button(btn_row, text="Clear", font=("Georgia", 9),
                  bg=self.BG_DARK, fg=self.FG_MUTED, activebackground=self.BG_MID,
                  relief="flat", cursor="hand2", padx=10, pady=4,
                  command=self._clear_docx).pack(side="left", padx=(8, 0))

        self.docx_count = tk.Label(btn_row, text="No files selected",
                                   font=("Courier", 9), bg=self.BG_DARK, fg=self.FG_MUTED)
        self.docx_count.pack(side="left", padx=14)

        self.file_list = tk.Listbox(section, font=("Courier", 8), bg=self.BG_PANEL,
                                    fg=self.FG_TEXT, relief="flat", height=6,
                                    selectbackground=self.BG_MID, bd=0)
        self.file_list.pack(fill="x", padx=8, pady=4)

        tk.Frame(self, bg=self.BG_MID, height=1).pack(fill="x", padx=30, pady=12)

    def _build_xlsx_section(self):
        section = tk.Frame(self, bg=self.BG_DARK)
        section.pack(fill="x", padx=30)

        tk.Label(section, text="2.  Excel Database (.xlsx)",
                 font=("Georgia", 10, "bold"), bg=self.BG_DARK, fg=self.FG_GOLD
                 ).pack(anchor="w")

        inner = tk.Frame(section, bg=self.BG_DARK)
        inner.pack(fill="x", pady=(6, 0))

        self.xlsx_var = tk.StringVar(value="No file selected")
        tk.Label(inner, textvariable=self.xlsx_var, font=("Courier", 8),
                 bg=self.BG_PANEL, fg=self.FG_MUTED, anchor="w", padx=8, pady=6,
                 relief="flat").pack(side="left", fill="x", expand=True)
        tk.Button(inner, text="Browse", font=("Georgia", 9),
                  bg=self.BG_MID, fg=self.FG_GOLD, activebackground="#2a4f7f",
                  relief="flat", cursor="hand2", padx=12,
                  command=self._pick_xlsx).pack(side="left", padx=(6, 0))

        tk.Frame(self, bg=self.BG_MID, height=1).pack(fill="x", padx=30, pady=14)

    def _build_action_buttons(self):
        frame = tk.Frame(self, bg=self.BG_DARK)
        frame.pack(padx=30, fill="x")

        self.btn = tk.Button(frame, text="Run Sync + Refresh Dashboard",
                             font=("Georgia", 12, "bold"),
                             bg=self.ACCENT, fg=self.BG_DARK, activebackground=self.FG_GOLD,
                             relief="flat", cursor="hand2", pady=10,
                             command=self._run)
        self.btn.pack(fill="x")

        self.btn_dash = tk.Button(frame,
                                  text="Refresh Dashboard Only  (no Word files needed)",
                                  font=("Georgia", 9),
                                  bg=self.BG_MID, fg=self.FG_GOLD, activebackground="#2a4f7f",
                                  relief="flat", cursor="hand2", pady=6,
                                  command=self._run_dash_only)
        self.btn_dash.pack(fill="x", pady=(6, 0))

    def _build_log_panel(self):
        tk.Label(self, text="Log", font=("Courier", 9),
                 bg=self.BG_DARK, fg=self.FG_MUTED).pack(anchor="w", padx=30, pady=(12, 2))

        log_frame = tk.Frame(self, bg=self.BG_DARK)
        log_frame.pack(padx=30, fill="both", expand=True, pady=(0, 20))

        self.log_box = tk.Text(log_frame, font=("Courier", 8),
                               bg=self.BG_PANEL, fg=self.FG_TEXT,
                               relief="flat", state="disabled", wrap="word",
                               bd=0, padx=10, pady=8)
        self.log_box.pack(fill="both", expand=True)

        self._log("Ready. Add Word files, select the Excel database, then click Run.")
        self._log("")
        self._log("   Simple single-block FDR cells are handled by the regex parser.")
        self._log("   Complex cells use the Anthropic API when ANTHROPIC_API_KEY is set,")
        self._log("   otherwise the regex parser is used as a fallback.")
        self._log("   Use 'Refresh Dashboard Only' to rebuild charts without a new sync.")

    # -- Event handlers ---------------------------------------------------

    def _add_docx(self):
        paths = filedialog.askopenfilenames(
            title="Select Word Report(s) - hold Ctrl for multiple",
            filetypes=[("Word Documents", "*.docx")],
        )
        for path in paths:
            if path not in self.docx_paths:
                self.docx_paths.append(path)
                self.file_list.insert("end", f"  {os.path.basename(path)}")
        n = len(self.docx_paths)
        self.docx_count.config(text=f"{n} file{'s' if n != 1 else ''} selected")

    def _clear_docx(self):
        self.docx_paths = []
        self.file_list.delete(0, "end")
        self.docx_count.config(text="No files selected")

    def _pick_xlsx(self):
        path = filedialog.askopenfilename(
            title="Select Excel Database",
            filetypes=[("Excel Files", "*.xlsx")],
        )
        if path:
            self.xlsx_path = path
            self.xlsx_var.set(os.path.basename(path))

    def _log(self, msg: str):
        self.log_box.configure(state="normal")
        self.log_box.insert("end", msg + "\n")
        self.log_box.see("end")
        self.log_box.configure(state="disabled")

    def _set_busy(self, busy: bool, label: str = "Run Sync + Refresh Dashboard"):
        state = "disabled" if busy else "normal"
        text  = "Processing..." if busy else label
        self.btn.configure(state=state, text=text)
        self.btn_dash.configure(state=state)

    def _clear_log(self):
        self.log_box.configure(state="normal")
        self.log_box.delete("1.0", "end")
        self.log_box.configure(state="disabled")

    def _run(self):
        if not self.docx_paths:
            messagebox.showwarning("No Word files", "Please add at least one Word document.")
            return
        if not self.xlsx_path:
            messagebox.showwarning("No Excel file", "Please select the Excel database.")
            return

        self._set_busy(True)
        self._clear_log()

        def task():
            try:
                sync_all(self.docx_paths, self.xlsx_path, self._log)
                self.after(0, lambda: messagebox.showinfo(
                    "Success",
                    f"All {len(self.docx_paths)} document(s) processed.\n"
                    "tbl_TER_Status, tbl_Metadata_Findings, tbl_CB_Needs updated.\n"
                    "Dashboard sheet rebuilt.\n\n"
                    "Upload the updated Excel back to SharePoint when ready.",
                ))
            except Exception as exc:
                self.after(0, lambda: messagebox.showerror("Error", str(exc)))
                self._log(f"\nError: {exc}")
            finally:
                self.after(0, lambda: self._set_busy(False))

        threading.Thread(target=task, daemon=True).start()

    def _run_dash_only(self):
        if not self.xlsx_path:
            messagebox.showwarning("No Excel file", "Please select the Excel database first.")
            return

        self._set_busy(True, "Refresh Dashboard Only")
        self._clear_log()

        def task():
            try:
                refresh_dashboard_only(self.xlsx_path, self._log)
                self.after(0, lambda: messagebox.showinfo(
                    "Dashboard refreshed",
                    "The Dashboard sheet has been rebuilt from current data.\n"
                    "Open the Excel file to see updated charts and KPIs.",
                ))
            except Exception as exc:
                self.after(0, lambda: messagebox.showerror("Error", str(exc)))
                self._log(f"\nError: {exc}")
            finally:
                self.after(0, lambda: self._set_busy(False))

        threading.Thread(target=task, daemon=True).start()


if __name__ == "__main__":
    App().mainloop()
